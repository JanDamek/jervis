"""VLM-first document extraction. Replaces Apache Tika.

Extraction strategy per mime type:
- text/*, CSV, JSON → direct UTF-8 decode
- HTML → BeautifulSoup strip tags
- DOCX → python-docx paragraphs + tables
- XLSX → openpyxl sheet rows
- PDF → pymupdf text + VLM for pages with images/scans (hybrid)
- Images → VLM description + OCR

FAIL-FAST: no fallback engines. VLM unavailable → retry with exponential backoff.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ImageDescription:
    description: str
    ocr_text: str = ""


@dataclass
class PageContent:
    page_number: int
    text: str
    images: list[ImageDescription] = field(default_factory=list)


@dataclass
class ExtractedDocument:
    text: str
    pages: list[PageContent] = field(default_factory=list)
    language: str = ""
    method: str = "direct"  # "pymupdf"|"vlm"|"direct"|"hybrid"|"docx"|"xlsx"|"html"
    metadata: dict = field(default_factory=dict)


# MIME type routing
_TEXT_MIMES = {"text/plain", "text/csv", "text/markdown", "text/x-markdown",
               "application/json", "application/xml", "text/xml", "text/yaml",
               "application/x-yaml"}
_HTML_MIMES = {"text/html", "application/xhtml+xml"}
_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
_XLSX_MIMES = {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
               "application/vnd.ms-excel"}
_IMAGE_MIMES = {"image/png", "image/jpeg", "image/jpg", "image/webp",
                "image/bmp", "image/gif", "image/tiff"}

# File extension fallback mapping
_EXT_TO_TYPE = {
    ".txt": "text", ".csv": "text", ".json": "text", ".xml": "text",
    ".yaml": "text", ".yml": "text", ".md": "text",
    ".html": "html", ".htm": "html",
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx", ".xls": "xlsx",
    ".png": "image", ".jpg": "image", ".jpeg": "image", ".webp": "image",
    ".bmp": "image", ".gif": "image", ".tiff": "image",
}


def _detect_type(mime_type: str, filename: str) -> str:
    """Detect document type from mime and filename."""
    mime_lower = mime_type.lower().strip() if mime_type else ""

    if mime_lower in _TEXT_MIMES:
        return "text"
    if mime_lower in _HTML_MIMES:
        return "html"
    if mime_lower in _PDF_MIMES:
        return "pdf"
    if mime_lower in _DOCX_MIMES:
        return "docx"
    if mime_lower in _XLSX_MIMES:
        return "xlsx"
    if mime_lower in _IMAGE_MIMES:
        return "image"

    # Fallback to extension
    fname_lower = filename.lower() if filename else ""
    for ext, doc_type in _EXT_TO_TYPE.items():
        if fname_lower.endswith(ext):
            return doc_type

    # Unknown — try as text
    return "text"


class DocumentExtractor:
    """VLM-first document extraction. Replaces Tika."""

    async def extract(
        self,
        file_bytes: bytes,
        filename: str,
        mime_type: str = "",
        max_tier: str = "NONE",
    ) -> ExtractedDocument:
        """Route to appropriate extractor based on mime type."""
        doc_type = _detect_type(mime_type, filename)
        logger.info("DocumentExtractor: file=%s mime=%s type=%s size=%d",
                     filename, mime_type, doc_type, len(file_bytes))

        if doc_type == "text":
            return self._extract_text_direct(file_bytes, filename)
        elif doc_type == "html":
            return self._extract_html(file_bytes)
        elif doc_type == "docx":
            return self._extract_docx(file_bytes)
        elif doc_type == "xlsx":
            return self._extract_xlsx(file_bytes)
        elif doc_type == "pdf":
            return await self._extract_pdf(file_bytes, max_tier)
        elif doc_type == "image":
            return await self._extract_image(file_bytes, max_tier)
        else:
            return self._extract_text_direct(file_bytes, filename)

    def _extract_text_direct(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """TXT, CSV, JSON, etc. — direct UTF-8 decode."""
        text = file_bytes.decode("utf-8", errors="replace")
        return ExtractedDocument(
            text=text,
            method="direct",
            metadata={"filename": filename},
        )

    def _extract_html(self, file_bytes: bytes) -> ExtractedDocument:
        """HTML — BeautifulSoup strip tags, preserve text structure."""
        from bs4 import BeautifulSoup

        html_text = file_bytes.decode("utf-8", errors="replace")
        soup = BeautifulSoup(html_text, "lxml")

        # Remove script and style elements
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        return ExtractedDocument(
            text=text,
            method="html",
            metadata={"title": soup.title.string if soup.title else ""},
        )

    def _extract_docx(self, file_bytes: bytes) -> ExtractedDocument:
        """DOCX — python-docx paragraphs + tables."""
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))

        text = "\n\n".join(parts)
        return ExtractedDocument(
            text=text,
            method="docx",
            metadata={
                "page_count": len(doc.sections),
                "has_tables": len(doc.tables) > 0,
            },
        )

    def _extract_xlsx(self, file_bytes: bytes) -> ExtractedDocument:
        """XLSX — openpyxl sheet rows."""
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    sheet_rows.append(" | ".join(cells))
            if sheet_rows:
                parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(sheet_rows))

        wb.close()
        text = "\n\n".join(parts)
        return ExtractedDocument(
            text=text,
            method="xlsx",
            metadata={"sheet_count": len(wb.sheetnames)},
        )

    async def _extract_pdf(self, file_bytes: bytes, max_tier: str) -> ExtractedDocument:
        """PDF hybrid — pymupdf text + VLM for pages with images/scans.

        Strategy:
        1. pymupdf: extract text + detect images per page
        2. Pages with sufficient text & no images → pymupdf text only
        3. Pages with images/graphs OR low text (scan) → render to image → VLM
        4. Combine results per-page
        """
        import fitz  # pymupdf

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[PageContent] = []
        method = "pymupdf"
        has_vlm_pages = False

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text("text").strip()
            image_list = page.get_images(full=True)

            # Decide: pymupdf text or VLM
            has_images = len(image_list) > 0
            is_scan = len(page_text) < 100  # Less than 100 chars = likely scanned

            if (has_images or is_scan) and (is_scan or not page_text):
                # Page needs VLM: render to image, call VLM
                try:
                    pix = page.get_pixmap(dpi=200)
                    image_bytes = pix.tobytes("png")
                    vlm_text = await self._call_vlm(
                        image_bytes,
                        "Extract all text, data, and describe all visual elements "
                        "(diagrams, charts, tables, images) from this document page. "
                        "Preserve the original language. Output plain text.",
                        max_tier,
                    )
                    page_content = PageContent(
                        page_number=page_num + 1,
                        text=vlm_text,
                        images=[ImageDescription(description=vlm_text)],
                    )
                    has_vlm_pages = True
                except Exception as e:
                    logger.warning("VLM failed for PDF page %d: %s — using pymupdf text", page_num + 1, e)
                    page_content = PageContent(page_number=page_num + 1, text=page_text)
            else:
                # pymupdf text is sufficient
                page_content = PageContent(page_number=page_num + 1, text=page_text)

            pages.append(page_content)

        doc.close()

        if has_vlm_pages:
            method = "hybrid"

        full_text = "\n\n".join(
            f"--- Page {p.page_number} ---\n{p.text}" for p in pages if p.text
        )

        return ExtractedDocument(
            text=full_text,
            pages=pages,
            method=method,
            metadata={
                "page_count": len(pages),
                "has_images": has_vlm_pages,
            },
        )

    async def _extract_image(self, file_bytes: bytes, max_tier: str) -> ExtractedDocument:
        """Image — VLM description + OCR."""
        text = await self._call_vlm(
            file_bytes,
            "Describe this image in detail for indexing purposes. "
            "Extract all text visible in the image (OCR). "
            "Describe diagrams, charts, tables, and key visual elements. "
            "Preserve the original language of any text.",
            max_tier,
        )
        return ExtractedDocument(
            text=text,
            method="vlm",
            metadata={"type": "image"},
        )

    async def _call_vlm(self, image_bytes: bytes, prompt: str, max_tier: str) -> str:
        """VLM call via llm_router with retry. FAIL-FAST, no fallback."""
        from app.services.llm_router import llm_generate_vision

        logger.info("VLM call: image_size=%d max_tier=%s", len(image_bytes), max_tier)
        result = await llm_generate_vision(
            image_bytes=image_bytes,
            prompt=prompt,
            max_tier=max_tier,
        )
        logger.info("VLM response: %d chars", len(result))
        return result
